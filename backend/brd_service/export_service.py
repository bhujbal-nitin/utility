"""
BRD Studio — Export Service
───────────────────────────
Generates AE-branded DOCX and PDF exports from persisted BRD sections.
Resolves [IMAGE_REF:capture_id] to embedded images.
Renders Mermaid diagrams as SVG/PNG.

CRITICAL BUG FIX: Reads from DB (latest persisted content), not frontend state.
"""

import os
import re
import io
import logging
import tempfile
import base64
import zlib
import json
import shutil
import subprocess
import textwrap
import httpx
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from xml.sax.saxutils import escape

from PIL import Image as PILImage

logger = logging.getLogger(__name__)

# AE brand colors
AE_NAVY = RGBColor(0x1F, 0x38, 0x64)
AE_ORANGE = RGBColor(0xE8, 0x70, 0x2A)
AE_DARK = RGBColor(0x0A, 0x16, 0x28)
AE_GRAY = RGBColor(0x8F, 0xA3, 0xC0)


class ExportService:
    """Export BRD sections to DOCX or PDF with AE branding."""

    def __init__(self, frames_dir: str, exports_dir: str):
        self.frames_dir = frames_dir
        self.exports_dir = exports_dir
        self._temp_files: List[str] = []

    def __del__(self):
        for path in self._temp_files:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass

    def _resolve_logo_path(self) -> Optional[str]:
        candidates = [
            os.path.join(os.path.dirname(__file__), "..", "assets", "AutomationEdge-Logo.png"),
            os.path.join(os.path.dirname(__file__), "..", "..", "AutomationEdge-Logo.png"),
            os.path.join(os.getcwd(), "AutomationEdge-Logo.png"),
        ]
        for c in candidates:
            if os.path.exists(c):
                return os.path.abspath(c)
        return None

    def _normalize_image_for_docx(self, img_path: str) -> io.BytesIO:
        """
        Convert any image file into a PNG BytesIO buffer that python-docx can reliably embed.
        FFmpeg-extracted JPEGs often have non-standard headers that python-docx's
        simple parser rejects with UnrecognizedImageError. PIL handles these gracefully.
        """
        img = PILImage.open(img_path)
        # Convert palette/RGBA modes to RGB for broad compatibility
        if img.mode in ("RGBA", "P", "LA"):
            background = PILImage.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            background.paste(img, mask=img.split()[-1] if "A" in img.mode else None)
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf

    def _find_capture_image(self, capture_id: str, captures: List[Dict]) -> Optional[str]:
        """Find the image file for a capture ID, preferring preview (user-edited) images."""
        target_id = str(capture_id).strip()
        
        # 1. Search in provided metadata list
        for cap in captures:
            cap_id = str(cap.get("id", ""))
            if cap_id != target_id:
                continue

            project_id = cap.get("project_id")
            frames_subdir = os.path.join(self.frames_dir, str(project_id)) if project_id else ""
            
            # Candidates to check
            candidate_paths = [
                cap.get("preview_image_path"),
                cap.get("image_path")
            ]
            
            # Add subdirectory candidates if project_id exists
            if frames_subdir:
                if cap.get("preview_image_path"):
                    candidate_paths.append(os.path.join(frames_subdir, os.path.basename(cap["preview_image_path"])))
                if cap.get("image_path"):
                    candidate_paths.append(os.path.join(frames_subdir, os.path.basename(cap["image_path"])))
                # Convention-based preview: preview_PID_CID.png or preview_CID.png
                candidate_paths.append(os.path.join(frames_subdir, f"preview_{target_id}.png"))

            for path in filter(None, candidate_paths):
                if os.path.exists(path) and os.path.getsize(path) > 0:
                    return os.path.abspath(path)

        # 2. Aggressive filesystem fallback: scan project directory if possible
        if captures:
            project_id = captures[0].get("project_id")
            if project_id:
                subdir = os.path.join(self.frames_dir, str(project_id))
                if os.path.isdir(subdir):
                    # Check for file containing target_id in its name
                    for root, dirs, files in os.walk(subdir):
                        for f in files:
                            if target_id in f:
                                return os.path.abspath(os.path.join(root, f))
        
        # 3. Global scan in frames_dir (last resort)
        if os.path.isdir(self.frames_dir):
            for root, dirs, files in os.walk(self.frames_dir):
                for f in files:
                    if target_id in f:
                        return os.path.abspath(os.path.join(root, f))

        # 4. Final fallback check: is the capture_id actually just the filename?
        if os.path.isfile(os.path.join(self.frames_dir, target_id)):
            return os.path.abspath(os.path.join(self.frames_dir, target_id))

        logger.warning(f"Could not resolve image for capture {capture_id}")
        return None

    # ── DOCX Export ───────────────────────────────────────────────────────

    async def export_docx(
        self,
        project_name: str,
        sections: List[Dict],
        captures: List[Dict],
        project_meta: Optional[Dict] = None,
    ) -> str:
        """
        Generate AE-branded DOCX from persisted sections.

        Args:
            project_name: Project name for filename
            sections: List of {section_key, title, content} from DB
            captures: List of capture dicts for IMAGE_REF resolution
            project_meta: Optional {client_name, ba_name, process_name}

        Returns:
            Absolute path to generated DOCX file
        """
        doc = Document()
        meta = project_meta or {}

        # Page setup
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        # Base font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = AE_DARK

        # Heading styles
        for i, (size, color) in enumerate([(18, AE_NAVY), (14, AE_NAVY), (12, AE_ORANGE)], start=1):
            hs = doc.styles[f"Heading {i}"]
            hs.font.name = "Calibri"
            hs.font.size = Pt(size)
            hs.font.color.rgb = color
            hs.font.bold = True

        # Page border
        self._apply_page_border(section)

        # Header with logo
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo = self._resolve_logo_path()
        if logo:
            header_para.add_run().add_picture(logo, width=Inches(1.2))

        # Footer
        footer = section.footer
        if footer.paragraphs:
            footer.paragraphs[0].text = ""
        footer_p = footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Confidential – AutomationEdge")
        run.font.size = Pt(8)
        run.font.color.rgb = AE_GRAY

        # ── Title page ────────────────────────────────────────────────────
        doc.add_paragraph("")  # Spacer
        doc.add_paragraph("")
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("Business Requirements Document")
        run.font.size = Pt(28)
        run.font.color.rgb = AE_NAVY
        run.font.bold = True
        run.font.name = "Calibri"

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(project_name)
        run.font.size = Pt(16)
        run.font.color.rgb = AE_ORANGE
        run.font.name = "Calibri"

        # Meta info
        doc.add_paragraph("")
        info_items = [
            ("Client", meta.get("client_name", "")),
            ("Process", meta.get("process_name", "")),
            ("Business Analyst", meta.get("ba_name", "")),
            ("Date", datetime.now().strftime("%B %d, %Y")),
        ]
        for label, value in info_items:
            if value:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"{label}: ")
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = AE_NAVY
                run2 = p.add_run(value)
                run2.font.size = Pt(11)

        doc.add_page_break()

        # ── Table of Contents placeholder ─────────────────────────────────
        doc.add_heading("Table of Contents", level=1)
        toc_p = doc.add_paragraph("(Update this field in Word: References → Update Table)")
        toc_p.italic = True
        doc.add_page_break()

        # ── Sections ──────────────────────────────────────────────────────
        for idx, sec in enumerate(sections):
            title = sec.get("title", f"Section {idx + 1}")
            content = sec.get("content", "")

            doc.add_heading(f"{idx + 1}. {title}", level=1)

            if not content.strip():
                doc.add_paragraph("No content generated for this section.")
                continue

            # Process content: resolve [IMAGE_REF] and render markdown
            self._render_content_to_docx(doc, content, captures)

            doc.add_paragraph("")  # Spacing between sections

        # Save
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", project_name).strip("._") or "Project"
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"BRD_{safe_name}_{timestamp}.docx"

        export_dir = os.path.join(self.exports_dir, safe_name)
        os.makedirs(export_dir, exist_ok=True)
        file_path = os.path.join(export_dir, filename)

        doc.save(file_path)
        logger.info(f"DOCX exported: {file_path}")
        return file_path

    async def export_docx_from_template(
        self,
        project_name: str,
        sections: List[Dict],
        captures: List[Dict],
        project_meta: Optional[Dict] = None,
    ) -> str:
        """
        Export sections to a DOCX using the BRD-Template.docx with [[tag]] replacement.
        """
        template_path = os.path.join(os.path.dirname(__file__), "templates", "BRD-Template.docx")
        if not os.path.exists(template_path):
            logger.error(f"Template not found at {template_path}, falling back to legacy export.")
            return await self.export_docx(project_name, sections, captures, project_meta)

        doc = Document(template_path)
        meta = project_meta or {}
        
        # 1. Build a lookup for sections
        section_by_key = {s["section_key"]: s["content"] for s in sections}
        
        # 2. Add numeric mappings to match template [[section_N_content]]
        numeric_mapping = {
            "section_1": "process_summary",
            "section_4": "applications_involved",
            "section_5": "feasibility_observations",
            "section_6": "io_details",
            "section_7": "flow_existing",
            "section_8": "flow_proposed",
            "section_9": "process_detail",
            "section_10": "validations",
            "section_11": "exceptions",
            "section_12": "rules",
            "section_13": "func_req",
            "section_14": "nonfunc_req",
            "section_15": "recommendations",
        }
        for tag, key in numeric_mapping.items():
            # Ensure the tag is ALWAYS initialized in section_by_key
            # Use data from internal key if present, otherwise empty string (will trigger "NA" replacement)
            section_by_key[tag] = section_by_key.get(key, "")

        # 3. Build metadata map for tags
        summary_map = {
            "ba_name": meta.get("ba_name", "AI Assistant"),
            "client_name": meta.get("client_name", "AutomationEdge"),
            "process_name": meta.get("process_name", project_name),
            "date": datetime.now().strftime("%B %d, %Y"),
            "project_name": project_name
        }

        # 4. Process all paragraphs and tables for [[tag]]
        async def process_para(p):
            txt = p.text or ""
            matches = re.findall(r"\[\[(.*?)\]\]", txt)
            for tag in matches:
                full_tag = f"[[{tag}]]"
                clean_tag = tag.strip().lower()
                
                # Case A: Bulky Section (Markdown content)
                # Map [[section_x_content]] or [[internal_key_content]]
                if "content" in clean_tag:
                    key = clean_tag.replace("_content", "")
                    content = section_by_key.get(key) or ""
                    
                    if content.strip():
                        # Clear the tag
                        p.text = p.text.replace(full_tag, "")
                        # Insert rendered blocks after this paragraph
                        await self._render_blocks_after_para(doc, p, content, captures)
                    else:
                        p.text = p.text.replace(full_tag, "NA")
                
                # Case B: Metadata tags
                elif clean_tag in summary_map:
                    p.text = p.text.replace(full_tag, summary_map[clean_tag])
                
                # Case C: Specialized tags
                elif clean_tag == "toc":
                    p.text = p.text.replace(full_tag, "")
                    self._add_toc_field(p)
                else:
                    # Check if internal key matches directly (no _content suffix)
                    if clean_tag in section_by_key:
                        content = section_by_key[clean_tag]
                        p.text = p.text.replace(full_tag, "")
                        await self._render_blocks_after_para(doc, p, content, captures)
                    else:
                        p.text = p.text.replace(full_tag, "-")

        for p in list(doc.paragraphs):
            await process_para(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in list(cell.paragraphs):
                        await process_para(p)

        # Save
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", project_name).strip("._") or "Project"
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"BRD_{safe_name}_{timestamp}.docx"
        
        export_dir = os.path.join(self.exports_dir, safe_name)
        os.makedirs(export_dir, exist_ok=True)
        file_path = os.path.join(export_dir, filename)
        
        doc.save(file_path)
        logger.info(f"Template-based DOCX exported: {file_path}")
        return file_path

    async def _render_blocks_after_para(self, doc: Document, p, content: str, captures: List[Dict]):
        """Render markdown blocks after a specific paragraph in template-based export."""
        insertion_point = p._p
        
        # Split content into blocks (para, list, table, image_ref, mermaid)
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Skip markdown horizontal rules.
            if re.match(r"^\s*[-*_]{3,}\s*$", line):
                i += 1
                continue

            # 1. Mermaid Rendering (fenced ```mermaid blocks)
            if "```mermaid" in line:
                mermaid_code = ""
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    mermaid_code += lines[i] + "\n"
                    i += 1
                if i < len(lines):
                    i += 1  # skip closing ```

                try:
                    img_bytes = await self._render_mermaid(mermaid_code)
                    if img_bytes:
                        new_p = doc.add_paragraph()
                        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = new_p.add_run()
                        run.add_picture(io.BytesIO(img_bytes), width=Inches(5.2))
                        insertion_point.addnext(new_p._p)
                        insertion_point = new_p._p
                except Exception as e:
                    logger.error(f"Mermaid render error: {e}")
                continue

            # 1b. Bare Mermaid (graph TD / flowchart / sequenceDiagram at line start)
            if re.match(r"^(graph\s+(TD|TB|LR|RL|BT)|flowchart\s+(TD|TB|LR|RL|BT)|sequenceDiagram)", line):
                mermaid_code = line + "\n"
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if not next_line or next_line.startswith("```"):
                        break
                    # Collect lines that look like Mermaid syntax (arrows, node defs, etc.)
                    if re.search(r"(-->|---|-\.-|==>|\|[^|]+\||subgraph|end\b)", next_line) or next_line.startswith("    "):
                        mermaid_code += lines[i] + "\n"
                        i += 1
                    else:
                        break

                try:
                    img_bytes = await self._render_mermaid(mermaid_code)
                    if img_bytes:
                        new_p = doc.add_paragraph()
                        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = new_p.add_run()
                        run.add_picture(io.BytesIO(img_bytes), width=Inches(5.2))
                        insertion_point.addnext(new_p._p)
                        insertion_point = new_p._p
                except Exception as e:
                    logger.error(f"Mermaid render error (bare): {e}")
                continue

            # 2. [IMAGE_REF] Resolution (handles mid-line or multiple)
            if "[IMAGE_REF" in line:
                # Support variant: [IMAGE_REF:id] or [IMAGE_REF: id]
                parts = re.split(r"(\[IMAGE_REF:?\s*[^\]\s]+\])", line)
                for part in parts:
                    img_match = re.search(r"\[IMAGE_REF:?\s*([^\]\s]+)\]", part)
                    if img_match:
                        cap_id = img_match.group(1).strip()
                        img_path = self._find_capture_image(cap_id, captures)
                        if img_path and os.path.exists(img_path) and os.path.getsize(img_path) > 0:
                            try:
                                img_buf = self._normalize_image_for_docx(img_path)
                                new_p = doc.add_paragraph()
                                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # Standardize to 6.0" width for professional boardroom reports
                                new_p.add_run().add_picture(img_buf, width=Inches(6.0))
                                insertion_point.addnext(new_p._p)
                                insertion_point = new_p._p
                                
                                cap = next((c for c in captures if c.get("id") == cap_id), None)
                                if cap:
                                    cap_p = doc.add_paragraph()
                                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = cap_p.add_run(f"Figure: {cap.get('label', 'Screenshot')}")
                                    run.font.italic = True
                                    run.font.size = Pt(9)
                                    run.font.color.rgb = AE_GRAY
                                    insertion_point.addnext(cap_p._p)
                                    insertion_point = cap_p._p
                            except Exception as e:
                                logger.error(f"Failed to embed capture {cap_id}: {type(e).__name__}: {e}", exc_info=True)
                                new_p = doc.add_paragraph(f"[Image: {cap_id} - Error embedding]")
                                insertion_point.addnext(new_p._p)
                                insertion_point = new_p._p
                        else:
                            logger.warning(f"Capture image not found or empty: {cap_id} (resolved path: {img_path})")
                    elif part.strip():
                        # Render text part
                        new_p = doc.add_paragraph()
                        self._add_run_with_markdown(new_p, part.strip())
                        insertion_point.addnext(new_p._p)
                        insertion_point = new_p._p
                i += 1
                continue

            # 3. Tables (accept imperfect markdown tables too)
            if self._is_table_candidate_line(line):
                table_lines = [line]
                i += 1
                while i < len(lines) and self._is_table_candidate_line(lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                temp_doc = Document()
                self._render_table_to_docx(temp_doc, table_lines)
                if temp_doc.tables:
                    table_elm = temp_doc.tables[0]._tbl
                    insertion_point.addnext(table_elm)
                    insertion_point = table_elm
                continue

            # 4. Standard Paragraph / Headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = min(4, max(1, len(heading_match.group(1))))
                heading_text = heading_match.group(2).strip()
                h = doc.add_heading(heading_text, level=level)
                insertion_point.addnext(h._p)
                insertion_point = h._p
                i += 1
                continue

            new_p = doc.add_paragraph()
            if line.startswith("- ") or line.startswith("* "):
                try:
                    new_p.style = "List Bullet"
                except KeyError:
                    new_p.add_run("• ")
                self._add_run_with_markdown(new_p, line[2:].strip())
            else:
                self._add_run_with_markdown(new_p, line)
            
            insertion_point.addnext(new_p._p)
            insertion_point = new_p._p
            i += 1

    def _add_run_with_markdown(self, paragraph, text):
        """Simple inline markdown parser for bold (**) and italic (* / _)"""
        text = text or ""
        # Normalize markdown links for DOCX output.
        text = re.sub(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", r"\1 (\2)", text)

        # Split by bold markers
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                # Bold
                content = part[2:-2]
                run = paragraph.add_run(content)
                run.font.bold = True
            else:
                # Normal or Italic
                inner_parts = re.split(r"(\*.*?\*|__.*?__|_.*?_)", part)
                for i_part in inner_parts:
                    if (i_part.startswith("*") and i_part.endswith("*")) or (i_part.startswith("_") and i_part.endswith("_")):
                        content = i_part[1:-1]
                        run = paragraph.add_run(content)
                        run.font.italic = True
                    elif i_part.startswith("__") and i_part.endswith("__"):
                        content = i_part[2:-2]
                        run = paragraph.add_run(content)
                        run.font.italic = True
                    else:
                        paragraph.add_run(i_part)

    async def _render_mermaid(self, code: str) -> Optional[bytes]:
        """Convert Mermaid to PNG with hardened sanitization and fallback chain."""
        try:
            code = self._sanitize_mermaid_code(code)
            code_bytes = code.encode("utf-8")

            # Primary: mermaid.ink (pako + base64)
            compressed = zlib.compress(code_bytes, 9)
            pako_b64 = base64.urlsafe_b64encode(compressed).decode("utf-8")
            pako_url = f"https://mermaid.ink/img/pako:{pako_b64}"
            base64_str = base64.b64encode(code_bytes).decode("utf-8")
            base64_url = f"https://mermaid.ink/img/{base64_str}"

            async with httpx.AsyncClient(follow_redirects=True) as client:
                for url in [pako_url, base64_url]:
                    try:
                        res = await client.get(url, timeout=25.0)
                        if res.status_code == 200 and res.content:
                            return res.content
                        logger.warning("Mermaid remote render returned %s", res.status_code)
                    except Exception as req_err:
                        logger.warning("Mermaid remote render failed: %s", req_err)

            # Fallback 1: Graphviz (if dot is installed)
            graphviz_png = self._render_graphviz_from_mermaid(code)
            if graphviz_png:
                logger.info("Diagram rendered via Graphviz fallback.")
                return graphviz_png

            # Fallback 2: deterministic textual flow card image
            text_png = self._render_text_flow_fallback(code)
            if text_png:
                logger.info("Diagram rendered via text-image fallback.")
                return text_png

            logger.error("All diagram render fallbacks failed.")
        except Exception as e:
            logger.error(f"Failed to render Mermaid diagram: {e}", exc_info=True)
        return None

    def _sanitize_mermaid_code(self, code: str) -> str:
        code = re.sub(r"```(mermaid)?", "", code, flags=re.IGNORECASE).strip()
        lines = []
        for raw in code.split("\n"):
            ln = raw.rstrip(";").rstrip()
            if not ln:
                continue
            # Remove very long labels that cause overflow
            ln = re.sub(
                r'\["([^"]{70,})"\]',
                lambda m: '["' + (m.group(1)[:67] + "...").replace('"', "'") + '"]',
                ln,
            )
            # Force quote any unquoted node label with []
            ln = re.sub(r'(\b[A-Za-z0-9_]+)\[([^\]"]+)\]', lambda m: f'{m.group(1)}["{m.group(2).strip().replace(chr(34), chr(39))}"]', ln)
            ln = re.sub(r'(\b[A-Za-z0-9_]+)\{([^\}"]+)\}', lambda m: f'{m.group(1)}{{"{m.group(2).strip().replace(chr(34), chr(39))}"}}', ln)
            lines.append(ln)
        cleaned = "\n".join(lines)
        if not re.match(r"^(graph|flowchart|sequenceDiagram|classDiagram|stateDiagram|erDiagram)\b", cleaned):
            cleaned = "graph TD\n" + cleaned
        return cleaned

    def _extract_edges_from_mermaid(self, code: str) -> List[tuple]:
        edges = []
        for ln in code.splitlines():
            line = ln.strip()
            if "-->" not in line:
                continue
            m = re.match(r"([A-Za-z0-9_]+)\s*--(?:\|[^|]+\|)?>\s*([A-Za-z0-9_]+)", line)
            if m:
                edges.append((m.group(1), m.group(2)))
        return edges

    def _render_graphviz_from_mermaid(self, code: str) -> Optional[bytes]:
        dot_exe = shutil.which("dot")
        if not dot_exe:
            return None
        edges = self._extract_edges_from_mermaid(code)
        if not edges:
            return None

        nodes = sorted({n for e in edges for n in e})
        dot_lines = ['digraph G {', 'rankdir=TB;', 'node [shape=box,fontsize=10];']
        for n in nodes:
            dot_lines.append(f'{n} [label="{n}"];')
        for src, dst in edges:
            dot_lines.append(f"{src} -> {dst};")
        dot_lines.append("}")
        dot_src = "\n".join(dot_lines)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".dot", mode="w", encoding="utf-8") as f:
            f.write(dot_src)
            dot_path = f.name
        out_path = dot_path + ".png"
        try:
            proc = subprocess.run([dot_exe, "-Tpng", dot_path, "-o", out_path], capture_output=True, text=True, check=False)
            if proc.returncode != 0 or not os.path.exists(out_path):
                return None
            with open(out_path, "rb") as f:
                return f.read()
        except Exception:
            return None
        finally:
            for p in [dot_path, out_path]:
                try:
                    if os.path.exists(p):
                        os.remove(p)
                except Exception:
                    pass

    def _render_text_flow_fallback(self, code: str) -> Optional[bytes]:
        try:
            # Create deterministic PNG with wrapped flow text.
            lines = [ln.strip() for ln in code.splitlines() if ln.strip()][:40]
            if not lines:
                lines = ["Flow diagram unavailable."]
            content = "\n".join(lines)
            wrapped = []
            for para in content.split("\n"):
                wrapped.extend(textwrap.wrap(para, width=68) or [""])
            width = 1200
            line_h = 24
            height = max(220, 80 + len(wrapped) * line_h)

            img = PILImage.new("RGB", (width, height), (248, 249, 251))
            from PIL import ImageDraw, ImageFont

            draw = ImageDraw.Draw(img)
            try:
                font = ImageFont.truetype("arial.ttf", 18)
                font_b = ImageFont.truetype("arialbd.ttf", 22)
            except Exception:
                font = ImageFont.load_default()
                font_b = ImageFont.load_default()

            draw.rectangle([(20, 20), (width - 20, height - 20)], outline=(31, 56, 100), width=2)
            draw.text((40, 36), "Process Flow (Fallback Rendering)", fill=(31, 56, 100), font=font_b)
            y = 78
            for ln in wrapped:
                draw.text((40, y), ln, fill=(30, 30, 30), font=font)
                y += line_h

            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return buf.getvalue()
        except Exception:
            return None

    def _add_toc_field(self, paragraph):
        """Native Word TOC field."""
        run = paragraph.add_run("[Table of Contents - Right-click to Update]")
        run.font.italic = True
        run.font.color.rgb = AE_GRAY
        
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        paragraph._p.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        paragraph._p.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        paragraph._p.append(fldChar2)

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        paragraph._p.append(fldChar3)

    def _render_content_to_docx(self, doc: Document, content: str, captures: List[Dict]):
        """Render markdown-ish content into DOCX paragraphs, tables, and images."""
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Handle [IMAGE_REF:capture_id] (mid-line or multiple)
            if "[IMAGE_REF" in line:
                parts = re.split(r"(\[IMAGE_REF:?\s*[^\]\s]+\])", line)
                for part in parts:
                    image_match = re.search(r"\[IMAGE_REF:?\s*([^\]\s]+)\]", part)
                    if image_match:
                        cap_id = image_match.group(1).strip()
                        img_path = self._find_capture_image(cap_id, captures)
                        if img_path:
                            try:
                                img_buf = self._normalize_image_for_docx(img_path)
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.add_run().add_picture(img_buf, width=Inches(6.0))
                                cap = next((c for c in captures if c.get("id") == cap_id), None)
                                if cap:
                                    cap_p = doc.add_paragraph()
                                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = cap_p.add_run(f"Figure: {cap.get('label', 'Screenshot')}")
                                    run.font.size = Pt(9)
                                    run.font.italic = True
                                    run.font.color.rgb = AE_GRAY
                            except Exception as e:
                                doc.add_paragraph(f"[Image: {cap_id}]")
                                logger.error(f"Failed to embed image {cap_id}: {type(e).__name__}: {e}", exc_info=True)
                        else:
                            doc.add_paragraph(f"[Image: {cap_id} not found]")
                    elif part.strip():
                        doc.add_paragraph(part.strip())
                i += 1
                continue

            # Handle Mermaid code blocks (render as text for now)
            if line.strip().startswith("```mermaid"):
                mermaid_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    mermaid_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```

                # Add as code block
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("[Flow Diagram - Mermaid]")
                run.font.italic = True
                run.font.color.rgb = AE_ORANGE
                run.font.size = Pt(10)

                for ml in mermaid_lines:
                    p = doc.add_paragraph()
                    run = p.add_run(ml)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                continue

            # Handle markdown tables
            if self._is_table_candidate_line(line):
                table_lines = [line]
                i += 1
                while i < len(lines) and self._is_table_candidate_line(lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                self._render_table_to_docx(doc, table_lines)
                continue

            # Handle headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
            if heading_match:
                level = min(4, max(1, len(heading_match.group(1))))
                doc.add_heading(heading_match.group(2).strip(), level=level)
                i += 1
                continue

            # Handle bullet points
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:]
                p = doc.add_paragraph(text, style="List Bullet")
                i += 1
                continue

            # Handle numbered items
            num_match = re.match(r"^\s*(\d+)\.\s+(.*)", line)
            if num_match:
                p = doc.add_paragraph(num_match.group(2), style="List Number")
                i += 1
                continue

            # Regular paragraph
            if line.strip():
                p = doc.add_paragraph()
                # Handle bold markers
                parts = re.split(r"(\*\*[^*]+\*\*)", line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    else:
                        p.add_run(part)

            i += 1

    def _is_table_candidate_line(self, line: str) -> bool:
        stripped = (line or "").strip()
        if not stripped:
            return False
        # Accept markdown separator rows and rows with at least 2 pipe delimiters.
        if re.match(r"^\|?[\s:-]+\|[\s|:-]*\|?$", stripped):
            return True
        return stripped.count("|") >= 2

    def _render_table_to_docx(self, doc: Document, table_lines: List[str]):
        """Render a markdown table into a DOCX table with AE styling."""
        rows = []
        for line in table_lines:
            if re.match(r"^\s*\|?[\s:-]+\|[\s|:-]*\|?\s*$", line):
                continue  # Skip separator row
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cells) < 2:
                continue
            rows.append(cells)

        if len(rows) < 2:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= num_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                self._add_run_with_markdown(p, cell_text.replace("`", ""))

                # Header row styling
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    self._set_cell_shading(cell, "1F3864")  # AE Navy
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    def _set_cell_shading(self, cell, color_hex: str):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def _apply_page_border(self, section):
        """Apply AE-branded page border."""
        try:
            sectPr = section._sectPr
            pgBorders = OxmlElement("w:pgBorders")
            pgBorders.set(qn("w:offsetFrom"), "page")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "24")
                border.set(qn("w:color"), "1F3864")
                pgBorders.append(border)
            sectPr.append(pgBorders)
        except Exception as e:
            logger.warning(f"Failed to apply page border: {e}")

    # ── PDF Export ────────────────────────────────────────────────────────

    def _convert_docx_to_pdf_best_effort(self, docx_path: str, pdf_path: str) -> bool:
        """
        Best-effort DOCX -> PDF conversion.
        Order:
        1) docx2pdf (Windows/macOS Word-backed conversion)
        2) LibreOffice/soffice headless conversion
        """
        # 1) docx2pdf python path
        try:
            from docx2pdf import convert as docx2pdf_convert  # type: ignore
            docx2pdf_convert(docx_path, pdf_path)
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return True
        except Exception:
            pass

        # 2) soffice headless path
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            try:
                out_dir = os.path.dirname(pdf_path)
                proc = subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        out_dir,
                        docx_path,
                    ],
                    capture_output=True,
                    text=True,
                    check=False,
                )
                expected = os.path.join(
                    out_dir,
                    f"{Path(docx_path).stem}.pdf",
                )
                if os.path.exists(expected) and os.path.getsize(expected) > 0:
                    if expected != pdf_path:
                        shutil.move(expected, pdf_path)
                    return True
                logger.warning("LibreOffice conversion failed rc=%s stderr=%s", proc.returncode, proc.stderr[:500])
            except Exception:
                pass
        return False

    async def export_pdf(
        self,
        project_name: str,
        sections: List[Dict],
        captures: List[Dict],
        project_meta: Optional[Dict] = None,
    ) -> str:
        """
        Generate PDF export.
        Canonical path:
          1) materialize template-mapped DOCX
          2) convert DOCX->PDF (if converter available)
          3) fallback to reportlab renderer
        """
        # Canonical materialization first.
        try:
            docx_path = await self.export_docx_from_template(project_name, sections, captures, project_meta)
            pdf_candidate = os.path.splitext(docx_path)[0] + ".pdf"
            if self._convert_docx_to_pdf_best_effort(docx_path, pdf_candidate):
                logger.info("PDF exported via canonical DOCX conversion: %s", pdf_candidate)
                return pdf_candidate
            logger.warning("DOCX->PDF converter unavailable; using reportlab fallback.")
        except Exception as e:
            logger.warning("Canonical PDF path failed, using fallback renderer: %s", e)

        # Fallback renderer
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", project_name).strip("._") or "Project"
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"BRD_{safe_name}_{timestamp}.pdf"

        export_dir = os.path.join(self.exports_dir, safe_name)
        os.makedirs(export_dir, exist_ok=True)
        file_path = os.path.join(export_dir, filename)

        doc = SimpleDocTemplate(file_path, pagesize=letter, topMargin=50, bottomMargin=50)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle("AETitle", parent=styles["Title"], fontSize=24, textColor=colors.HexColor("#1F3864"))
        heading_style = ParagraphStyle("AEHeading", parent=styles["Heading1"], fontSize=14, textColor=colors.HexColor("#1F3864"))
        body_style = ParagraphStyle("AEBody", parent=styles["Normal"], fontSize=10, leading=14)

        elements = []

        # Title
        elements.append(Spacer(1, 100))
        elements.append(Paragraph(escape("Business Requirements Document"), title_style))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph(escape(project_name), ParagraphStyle(
            "AESubtitle", parent=styles["Normal"], fontSize=16, textColor=colors.HexColor("#E8702A")
        )))
        elements.append(Spacer(1, 40))

        meta = project_meta or {}
        for label, value in [("Client", meta.get("client_name", "")), ("Date", datetime.now().strftime("%B %d, %Y"))]:
            if value:
                elements.append(Paragraph(f"<b>{escape(label)}:</b> {escape(value)}", body_style))

        elements.append(Spacer(1, 50))

        # Sections
        for idx, sec in enumerate(sections):
            title = sec.get("title", f"Section {idx + 1}")
            content = sec.get("content", "")

            elements.append(Paragraph(f"{idx + 1}. {escape(title)}", heading_style))
            elements.append(Spacer(1, 10))

            if content.strip():
                # Handle [IMAGE_REF] in content for PDF
                clean_content = re.sub(r"\[IMAGE_REF:[^\]]+\]", "[Screenshot]", content)
                # Split by paragraphs
                for para in clean_content.split("\n"):
                    para = para.strip()
                    if para:
                        try:
                            elements.append(Paragraph(escape(para), body_style))
                            elements.append(Spacer(1, 4))
                        except Exception:
                            pass  # Skip unparseable paragraphs
            else:
                elements.append(Paragraph("No content generated.", body_style))

            elements.append(Spacer(1, 20))

        doc.build(elements)
        logger.info(f"PDF exported: {file_path}")
        return file_path
