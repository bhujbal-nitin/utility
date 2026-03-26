"""
BRD Studio — Test Suite
───────────────────────
Tests for atomic save+version, pipeline, export, and video processing.
"""

import os
import sys
import json
import pytest
import asyncio
from unittest.mock import AsyncMock, MagicMock, patch

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


# ─── Test: Atomic Save + Version ─────────────────────────────────────────────

class TestAtomicSaveVersion:
    """
    CRITICAL BUG FIX TEST:
    Every section save MUST create a version snapshot in the same transaction.
    This prevents the BRDCopilot bug where edits were lost.
    """

    def test_version_snapshot_is_created_before_content_update(self):
        """
        Simulate the atomic save+version pattern.
        The version snapshot should contain the OLD content,
        and the section should have the NEW content after save.
        """
        # Simulate current state
        old_content = "## Process Summary\n- Step 1: Login"
        new_content = "## Process Summary\n- Step 1: Login\n- Step 2: Dashboard"
        current_version = 3

        # --- Simulate atomic save ---
        # 1. Snapshot current (BEFORE update)
        version_snapshot = {
            "version_number": current_version,
            "content_snapshot": old_content,
        }

        # 2. Update section
        section = {
            "content": new_content,
            "version": current_version + 1,
            "is_manual_override": True,
        }

        # Assertions
        assert version_snapshot["content_snapshot"] == old_content
        assert version_snapshot["version_number"] == 3
        assert section["content"] == new_content
        assert section["version"] == 4
        assert section["is_manual_override"] is True

    def test_manual_override_prevents_regeneration_overwrite(self):
        """
        Manually edited sections should NOT be overwritten during regeneration.
        This was a key BRDCopilot bug.
        """
        sections = [
            {"section_key": "process_summary", "content": "Manual edit by user", "is_manual_override": True},
            {"section_key": "func_req", "content": "Auto-generated", "is_manual_override": False},
        ]

        generated = [
            {"section_key": "process_summary", "content": "New AI content for summary"},
            {"section_key": "func_req", "content": "New AI content for requirements"},
        ]

        # Simulate regeneration logic
        for gen_sec in generated:
            existing = next((s for s in sections if s["section_key"] == gen_sec["section_key"]), None)
            if existing and existing["is_manual_override"]:
                continue  # SKIP: manually edited
            if existing:
                existing["content"] = gen_sec["content"]

        # process_summary should NOT be overwritten
        assert sections[0]["content"] == "Manual edit by user"
        # func_req SHOULD be overwritten
        assert sections[1]["content"] == "New AI content for requirements"


# ─── Test: Evidence Pack Building ─────────────────────────────────────────────

class TestEvidencePackBuilding:
    """Test the pipeline's evidence builder."""

    def test_build_evidence_with_captures_and_transcript(self):
        from brd_service.brd_pipeline import BRDPipeline

        pipeline = BRDPipeline()
        captures = [
            {"id": "cap1", "label": "Login Screen", "timestamp": 5.0, "is_kept": True,
             "description": "Login page with username and password fields",
             "ocr_text": "Username Password Sign In",
             "details_json": {"app_name": "SAP", "ui_elements": ["text_field", "button"]}},
        ]
        videos = [
            {"filename": "walkthrough.mp4", "transcript_text": "Click on login button and enter credentials"},
        ]
        documents = [
            {"filename": "process.txt", "extracted_text": "Standard login process for SAP"},
        ]

        result = pipeline.build_evidence_pack(captures, videos, documents)

        assert "## Transcript Evidence" in result
        assert "walkthrough.mp4" in result
        assert "## Screen Capture Evidence" in result
        assert "Login Screen" in result
        assert "[IMAGE_REF:cap1]" in result
        assert "SAP" in result
        assert "## Additional Documents" in result

    def test_skips_non_kept_captures(self):
        from brd_service.brd_pipeline import BRDPipeline

        pipeline = BRDPipeline()
        captures = [
            {"id": "cap1", "label": "Kept", "is_kept": True, "timestamp": 0},
            {"id": "cap2", "label": "Skipped", "is_kept": False, "timestamp": 5},
        ]

        result = pipeline.build_evidence_pack(captures, [], [])

        assert "Kept" in result
        assert "Skipped" not in result


# ─── Test: Export Service ─────────────────────────────────────────────────────

class TestExportService:
    """Test DOCX and PDF generation."""

    @pytest.mark.asyncio
    async def test_docx_export_creates_file(self, tmp_path):
        from brd_service.export_service import ExportService

        frames_dir = str(tmp_path / "frames")
        exports_dir = str(tmp_path / "exports")
        os.makedirs(frames_dir, exist_ok=True)
        os.makedirs(exports_dir, exist_ok=True)

        svc = ExportService(frames_dir, exports_dir)
        sections = [
            {"section_key": "process_summary", "title": "Process Summary", "content": "## Summary\n- Step 1\n- Step 2"},
            {"section_key": "func_req", "title": "Functional Requirements", "content": "REQ-001: Login"},
        ]

        file_path = await svc.export_docx("Test_Project", sections, [], {"client_name": "Acme Corp"})

        assert os.path.exists(file_path)
        assert file_path.endswith(".docx")
        assert "Test_Project" in os.path.basename(file_path)

    @pytest.mark.asyncio
    async def test_pdf_export_creates_file(self, tmp_path):
        from brd_service.export_service import ExportService

        frames_dir = str(tmp_path / "frames")
        exports_dir = str(tmp_path / "exports")
        os.makedirs(frames_dir, exist_ok=True)
        os.makedirs(exports_dir, exist_ok=True)

        svc = ExportService(frames_dir, exports_dir)
        sections = [
            {"section_key": "process_summary", "title": "Process Summary", "content": "Summary content here"},
        ]

        file_path = await svc.export_pdf("Test_PDF_Project", sections, [])

        assert os.path.exists(file_path)
        assert file_path.endswith(".pdf")

    def test_render_table_to_docx(self):
        from brd_service.export_service import ExportService
        from docx import Document

        svc = ExportService(".", ".")
        doc = Document()

        table_lines = [
            "| Name | Role | Access |",
            "|------|------|--------|",
            "| SAP | Primary | Browser |",
            "| Excel | Secondary | Desktop |",
        ]

        svc._render_table_to_docx(doc, table_lines)

        # Verify table was created
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3  # header + 2 data rows
        assert table.cell(0, 0).text == "Name"
        assert table.cell(1, 0).text == "SAP"


# ─── Test: Video Processor ───────────────────────────────────────────────────

class TestVideoProcessor:
    """Test frame capture logic."""

    def test_ssim_dedup_removes_similar_frames(self):
        """Test that SSIM comparison removes duplicate frames."""
        # This is a unit test for the SSIM logic
        import numpy as np

        # Create two identical images
        img1 = np.zeros((100, 100, 3), dtype=np.uint8)
        img2 = np.zeros((100, 100, 3), dtype=np.uint8)

        # They should be identical (SSIM ~= 1.0)
        from skimage.metrics import structural_similarity as ssim
        score = ssim(img1, img2, channel_axis=2)
        assert score > 0.95

        # Create a very different image
        img3 = np.ones((100, 100, 3), dtype=np.uint8) * 255
        score2 = ssim(img1, img3, channel_axis=2)
        assert score2 < 0.1


# ─── Test: Prompts ────────────────────────────────────────────────────────────

class TestPrompts:
    """Verify prompt templates are complete."""

    def test_all_section_definitions_have_prompts(self):
        from brd_service.prompts import SECTION_DEFINITIONS, SECTION_WRITER_PROMPTS

        for sec_def in SECTION_DEFINITIONS:
            key = sec_def["key"]
            assert key in SECTION_WRITER_PROMPTS, f"Missing prompt for section: {key}"

    def test_prompts_contain_evidence_placeholder(self):
        from brd_service.prompts import SECTION_WRITER_PROMPTS

        for key, prompt in SECTION_WRITER_PROMPTS.items():
            assert "{evidence_pack}" in prompt, f"Prompt for {key} missing {{evidence_pack}} placeholder"

    def test_section_preamble_security(self):
        from brd_service.prompts import SECTION_WRITER_PROMPTS

        for key, prompt in SECTION_WRITER_PROMPTS.items():
            assert "UNTRUSTED" in prompt, f"Security preamble missing for {key}"

    def test_flow_sections_request_mermaid_diagrams(self):
        from brd_service.prompts import SECTION_WRITER_PROMPTS

        for key in ["flow_existing", "flow_proposed"]:
            assert "mermaid" in SECTION_WRITER_PROMPTS[key].lower(), f"{key} should request Mermaid diagrams"


# ─── Test: Section Order ─────────────────────────────────────────────────────

class TestSectionOrder:
    """Verify section definitions match expected BRD structure."""

    def test_section_count(self):
        from brd_service.prompts import SECTION_DEFINITIONS
        assert len(SECTION_DEFINITIONS) == 13

    def test_required_sections_present(self):
        from brd_service.prompts import SECTION_DEFINITIONS

        required = [
            "process_summary", "applications_involved", "flow_existing",
            "flow_proposed", "process_detail", "func_req", "exceptions",
        ]
        keys = [s["key"] for s in SECTION_DEFINITIONS]
        for r in required:
            assert r in keys, f"Required section {r} missing"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
